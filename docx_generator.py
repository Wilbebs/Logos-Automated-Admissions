import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from templates.report_template import get_report_structure

def generate_report(student_data, classification):
    """
    Main compatibility wrapper - calls the enhanced generator
    """
    generator = ComprehensiveReportGenerator()
    return generator.generate_comprehensive_report(student_data, classification)

class ComprehensiveReportGenerator:
    def __init__(self):
        self.doc = Document()
    
    def _setup_document_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

    def _add_horizontal_line(self, color='003366', width=12):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1
        
        # Add a border-bottom to simulate a horizontal line
        p_pr = p._p.get_or_add_pPr()
        p_borders = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(width))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color)
        p_borders.append(bottom)
        p_pr.append(p_borders)

    def _add_professional_header(self):
        """Enhanced header with institutional branding"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        title_run = title.add_run('UNIVERSIDAD CRISTIANA DE LOGOS')
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)  # UCL Blue
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle_run = subtitle.add_run('Reporte de Clasificación Académica')
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(204, 153, 0)  # UCL Gold
        
        self._add_horizontal_line(color='003366', width=12)
        self.doc.add_paragraph()

    def _add_executive_summary_box(self, student_data, classification):
        """Shaded box with key recommendation"""
        summary_para = self.doc.add_paragraph()
        summary_para.paragraph_format.space_before = Pt(12)
        summary_para.paragraph_format.space_after = Pt(12)
        
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E6F2FF')  # Light blue
        summary_para._p.get_or_add_pPr().append(shading_elm)
        
        summary_para.paragraph_format.left_indent = Inches(0.3)
        summary_para.paragraph_format.right_indent = Inches(0.3)
        
        header_run = summary_para.add_run('📋 RECOMENDACIÓN ACADÉMICA\n\n')
        header_run.font.bold = True
        header_run.font.size = Pt(14)
        
        summary_para.add_run('Nivel Recomendado: ').font.bold = True
        
        level_run = summary_para.add_run(classification['recommended_level'])
        level_run.font.size = Pt(18)
        level_run.font.bold = True
        level_run.font.color.rgb = RGBColor(0, 51, 102)
        
        summary_para.add_run(f'\n\nConfianza: {classification.get("confidence_score", "N/A")}%')
        summary_para.add_run(f'\nFecha: {datetime.now().strftime("%d de %B, %Y")}')
        
        self.doc.add_paragraph()

    def _add_applicant_info_section(self, student_data):
        self.doc.add_heading('Información del Solicitante', level=1)
        
        info_para = self.doc.add_paragraph()
        info_para.add_run('Nombre: ').font.bold = True
        info_para.add_run(f"{student_data.get('applicant_name', 'N/A')}\n")
        
        info_para.add_run('Email: ').font.bold = True
        info_para.add_run(f"{student_data.get('email', 'N/A')}\n")
        
        info_para.add_run('Interés: ').font.bold = True
        info_para.add_run(f"{student_data.get('program_interest', 'N/A')}\n")
        
        info_para.add_run('Nivel Educativo: ').font.bold = True
        info_para.add_run(f"{student_data.get('education_level', 'N/A')}")
        
        self.doc.add_paragraph()

    def _add_reasoning_section(self, classification):
        """Add detailed AI reasoning breakdown"""
        self.doc.add_heading('Análisis Detallado de Clasificación', level=1)
        
        reasoning = classification.get('reasoning', {})
        
        # Educational Assessment
        self.doc.add_heading('1. Evaluación Académica', level=2)
        self.doc.add_paragraph(reasoning.get('educational_assessment', 'N/A'))
        
        # Ministry Experience
        self.doc.add_heading('2. Experiencia Ministerial', level=2)
        self.doc.add_paragraph(reasoning.get('ministry_experience_assessment', 'N/A'))
        
        # Pastoral Recommendation
        self.doc.add_heading('3. Recomendación Pastoral', level=2)
        self.doc.add_paragraph(reasoning.get('pastoral_recommendation_assessment', 'N/A'))
        
        self.doc.add_paragraph()

    def _add_program_details_section(self, classification):
        """Add detailed program information in table format"""
        self.doc.add_heading('Programas Recomendados', level=1)
        
        try:
            from classification_framework import PROGRAM_DETAILS
        except ImportError:
            PROGRAM_DETAILS = {}
        
        for program_name in classification.get('recommended_programs', []):
            program_info = PROGRAM_DETAILS.get(program_name, {})
            
            self.doc.add_heading(program_name, level=2)
            
            table = self.doc.add_table(rows=7, cols=2)
            table.style = 'Light Grid Accent 1'
            
            details = [
                ('📅 Duración', program_info.get('duration', 'N/A')),
                ('📚 Créditos', program_info.get('credits', 'N/A')),
                ('💻 Formato', program_info.get('format', 'N/A')),
                ('🎯 Enfoque', program_info.get('focus', 'N/A')),
                ('📋 Prerequisitos', program_info.get('prerequisites', 'N/A')),
                ('👥 Ideal Para', ', '.join(program_info.get('best_for', ['N/A'])[:2])),
                ('💰 Costo Estimado', program_info.get('cost_estimate', 'Consultar con admisiones'))
            ]
            
            for i, (label, value) in enumerate(details):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
                table.rows[i].cells[1].text = value
            
            self.doc.add_paragraph()

    def _add_next_steps_checklist(self, classification):
        """Add actionable next steps with checkboxes"""
        self.doc.add_heading('📝 Próximos Pasos', level=1)
        
        self.doc.add_paragraph('Para completar su proceso de admisión, siga estos pasos en orden:')
        
        steps = classification.get('next_steps', [
            'Revisar los programas recomendados en detalle',
            'Preparar documentos adicionales requeridos',
            'Contactar a la oficina de admisiones'
        ])
        
        for i, step in enumerate(steps, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.add_run('☐ ').font.color.rgb = RGBColor(0, 51, 102)
            p.add_run(f'{i}. {step}')
        
        self.doc.add_paragraph()

    def _add_missing_documents_section(self, classification):
        """Add missing documents alert if documents are incomplete"""
        missing_docs = classification.get('reasoning', {}).get('documents_missing', [])
        
        if missing_docs:
            self.doc.add_heading('⚠️ Documentos Faltantes', level=1)
            
            alert_para = self.doc.add_paragraph()
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'FFF4E6')  # Light orange
            alert_para._p.get_or_add_pPr().append(shading_elm)
            
            alert_para.paragraph_format.left_indent = Inches(0.3)
            
            warning_text = alert_para.add_run('Su clasificación está pendiente de los siguientes documentos:\n\n')
            warning_text.font.bold = True
            warning_text.font.color.rgb = RGBColor(204, 102, 0)
            
            for doc in missing_docs:
                alert_para.add_run(f'• {doc}\n')
            
            alert_para.add_run('\nPor favor envíe estos documentos a admisiones@logos.edu.')
            
            self.doc.add_paragraph()

    def _add_pathway_section(self, classification):
        """Add educational pathway if applicant is over-aspiring"""
        pathway = classification.get('reasoning', {}).get('pathway_explanation')
        
        if pathway:
            self.doc.add_heading('🎓 Su Trayectoria Académica Recomendada', level=1)
            
            path_para = self.doc.add_paragraph()
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F0F8FF')  # Alice blue
            path_para._p.get_or_add_pPr().append(shading_elm)
            
            path_para.paragraph_format.left_indent = Inches(0.3)
            path_para.add_run(pathway)
            
            self.doc.add_paragraph()

    def _add_contact_footer(self):
        """Add professional footer with contact information"""
        self._add_horizontal_line(color='CCCCCC', width=6)
        
        contact_heading = self.doc.add_paragraph()
        contact_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = contact_heading.add_run('Oficina de Admisiones')
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(0, 51, 102)
        
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.add_run('📧 admisiones@logos.edu | 📞 +1 (XXX) XXX-XXXX | 🌐 www.logos.edu')
        
        disclaimer = self.doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disc_run = disclaimer.add_run(
            'Esta es una recomendación generada por nuestro sistema de clasificación basado en IA. '
            'La decisión final corresponde al equipo de admisiones.'
        )
        disc_run.font.size = Pt(8)
        disc_run.font.italic = True
        disc_run.font.color.rgb = RGBColor(128, 128, 128)

    def generate_comprehensive_report(self, student_data, classification):
        self._setup_document_styles()
        self._add_professional_header()
        self._add_executive_summary_box(student_data, classification)
        self._add_applicant_info_section(student_data)
        self._add_reasoning_section(classification)
        self._add_program_details_section(classification)
        self._add_pathway_section(classification)
        self._add_next_steps_checklist(classification)
        self._add_missing_documents_section(classification)
        self._add_contact_footer()
        
        from docx_generator import _generate_filename
        filename = _generate_filename(student_data.get('applicant_name', 'Student'))
        filepath = os.path.join('/tmp', filename)
        
        # Ensure /tmp exists (especially on local windows mock)
        if not os.path.exists('/tmp'):
            os.makedirs('/tmp')
            
        self.doc.save(filepath)
        return filepath

def _render_header(doc, header, metadata):
    title = doc.add_heading(header['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading(header['subtitle'], level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Fecha: {metadata['generated_date']}")
    date_run.bold = True
    
    doc.add_paragraph()

def _render_section(doc, section):
    doc.add_heading(section['title'], level=2)
    
    section_type = section['type']
    
    if section_type == 'key_value_pairs':
        _render_key_value_pairs(doc, section['content'])
    elif section_type == 'recommendation':
        _render_recommendation(doc, section['content'])
    elif section_type == 'paragraph':
        doc.add_paragraph(section['content'])
    elif section_type == 'subsections':
        _render_subsections(doc, section['content'])
    elif section_type == 'alert_box':
        _render_alert_box(doc, section['content'], section.get('alert_level', 'info'))
    
    doc.add_paragraph()

def _render_key_value_pairs(doc, pairs):
    for item in pairs:
        p = doc.add_paragraph()
        p.add_run(f"{item['label']}: ").bold = True
        p.add_run(item['value'])

def _render_recommendation(doc, content):
    level_data = content['level']
    level_para = doc.add_paragraph()
    level_para.add_run(f"{level_data['label']}: ").bold = True
    
    level_run = level_para.add_run(level_data['value'])
    level_run.bold = True
    level_run.font.size = Pt(14)
    
    if level_data.get('highlight') and level_data.get('color') == 'blue':
        level_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    
    programs_data = content['programs']
    programs_para = doc.add_paragraph()
    programs_para.add_run(f"{programs_data['label']}:").bold = True
    
    for program in programs_data['value']:
        doc.add_paragraph(f'• {program}', style='List Bullet')
    
    doc.add_paragraph()
    
    confidence_data = content['confidence']
    conf_para = doc.add_paragraph()
    conf_para.add_run(f"{confidence_data['label']}: ").bold = True
    conf_para.add_run(confidence_data['value'])
    
    if confidence_data.get('show_warning'):
        warning = doc.add_paragraph()
        warning_run = warning.add_run('⚠️ ATENCIÓN: Nivel de confianza bajo. Revisión manual recomendada.')
        warning_run.bold = True
        warning_run.font.color.rgb = RGBColor(204, 0, 0)

def _render_subsections(doc, subsections):
    for subsection in subsections:
        subtitle_para = doc.add_paragraph()
        subtitle_para.add_run(subsection['subtitle']).bold = True
        doc.add_paragraph(subsection['text'])
        doc.add_paragraph()

def _render_alert_box(doc, content, alert_level):
    alert_para = doc.add_paragraph(content)
    
    if alert_level == 'warning':
        alert_para.runs[0].font.color.rgb = RGBColor(204, 102, 0)
    elif alert_level == 'error':
        alert_para.runs[0].font.color.rgb = RGBColor(204, 0, 0)

def _render_footer(doc, footer):
    doc.add_paragraph()
    doc.add_paragraph('_' * 50)
    
    footer_para = doc.add_paragraph(footer['text'])
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.italic = True
    
    if 'disclaimer' in footer:
        disclaimer_para = doc.add_paragraph(footer['disclaimer'])
        disclaimer_para.runs[0].font.size = Pt(8)
        disclaimer_para.runs[0].font.italic = True
        disclaimer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

def _generate_filename(applicant_name):
    safe_name = applicant_name.replace(' ', '_').replace('/', '_')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"Clasificacion_{safe_name}_{timestamp}.docx"